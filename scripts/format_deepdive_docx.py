"""
format_deepdive_docx.py — convert a deep-dive .md report into a structured Word document.

Usage (called automatically after company_deep_report.py generates the report):
    python scripts/format_deepdive_docx.py <report.md> [--out <output.docx>]

Or import and call:
    from format_deepdive_docx import md_to_docx
    docx_bytes = md_to_docx(md_text, company_name, symbol, isin)
"""
from __future__ import annotations
import os, io, re, sys, argparse, datetime as dt
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")

# --------------------------------------------------------------------------
# Section patterns from comapnydeepdive_prompt.txt output
# --------------------------------------------------------------------------
SECTION_PATTERNS = [
    (r"LAYER\s+1.*?EXECUTIVE SCORECARD",        "Executive Scorecard"),
    (r"A\.\s*Source Coverage",                   "A. Source Coverage"),
    (r"B\.\s*Business Profile",                  "B. Business Profile"),
    (r"C\.\s*Capital Efficiency",                "C. Capital Efficiency"),
    (r"D\.\s*Forensic Findings",                 "D. Forensic Findings"),
    (r"E\.\s*Fraud.*?Governance",                "E. Fraud & Governance"),
    (r"F\.\s*Risk Scorecard",                    "F. Risk Scorecard"),
    (r"G\.\s*Investment Thesis",                 "G. Investment Thesis"),
    (r"H\.\s*Earnings Quality",                  "H. Earnings Quality"),
    (r"I\.\s*Data Reliability",                  "I. Data Reliability"),
    (r"LAYER\s+2.*?ANALYTICAL NOTES",            "Layer 2 — Analytical Notes"),
    (r"PM ONE[-\s]PAGER",                        "PM One-Pager"),
]


def _add_cover(doc: Document, name: str, symbol: str, isin: str, coverage: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEEP DIVE REPORT")
    run.bold = True; run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(name)
    r2.bold = True; r2.font.size = Pt(18)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"{symbol}  ·  {isin}")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f"Generated {dt.datetime.now():%Y-%m-%d %H:%M}")

    if coverage:
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.add_run(f"Coverage: {coverage}").italic = True

    doc.add_page_break()


def _add_toc(doc: Document):
    h = doc.add_heading("Table of Contents", level=1)
    toc_sections = [
        "Executive Scorecard",
        "  A. Source Coverage",
        "  B. Business Profile",
        "  C. Capital Efficiency",
        "  D. Forensic Findings",
        "  E. Fraud & Governance",
        "  F. Risk Scorecard",
        "  G. Investment Thesis",
        "  H. Earnings Quality",
        "  I. Data Reliability",
        "Layer 2 — Analytical Notes",
        "PM One-Pager",
    ]
    for s in toc_sections:
        doc.add_paragraph(s, style="List Bullet" if not s.startswith("  ") else "List Bullet 2")
    doc.add_page_break()


def _parse_table(lines: list[str], doc: Document):
    rows = [l for l in lines if l.strip().startswith("|") and "---" not in l]
    if not rows:
        return False
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return False
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
    doc.add_paragraph()
    return True


def _flush_block(doc: Document, current_heading: Optional[str], block_lines: list[str]):
    if not block_lines:
        return
    # detect table blocks
    table_lines = [l for l in block_lines if l.strip().startswith("|")]
    if len(table_lines) >= 2:
        non_table = [l for l in block_lines if not l.strip().startswith("|") and "---" not in l]
        for l in non_table:
            if l.strip():
                doc.add_paragraph(l.strip())
        _parse_table(table_lines, doc)
        return
    # normal paragraphs
    para_buf = []
    for line in block_lines:
        stripped = line.strip()
        if not stripped:
            if para_buf:
                doc.add_paragraph(" ".join(para_buf))
                para_buf = []
        elif stripped.startswith("- ") or stripped.startswith("* "):
            if para_buf:
                doc.add_paragraph(" ".join(para_buf)); para_buf = []
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            if para_buf:
                doc.add_paragraph(" ".join(para_buf)); para_buf = []
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        elif stripped.startswith("### "):
            if para_buf:
                doc.add_paragraph(" ".join(para_buf)); para_buf = []
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            if para_buf:
                doc.add_paragraph(" ".join(para_buf)); para_buf = []
            doc.add_heading(stripped[3:], level=2)
        else:
            para_buf.append(stripped)
    if para_buf:
        doc.add_paragraph(" ".join(para_buf))


def md_to_docx(md_text: str, name: str = "", symbol: str = "",
               isin: str = "", coverage: str = "") -> bytes:
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    _add_cover(doc, name or "Company", symbol or "", isin or "", coverage)
    _add_toc(doc)

    lines = md_text.splitlines()
    current_heading = None
    block: list[str] = []

    def is_section_header(line: str) -> Optional[str]:
        clean = re.sub(r"[=\-#*]+", "", line).strip()
        for pattern, label in SECTION_PATTERNS:
            if re.search(pattern, clean, re.IGNORECASE):
                return label
        # generic h1/h2 in output
        m = re.match(r"^#{1,2}\s+(.+)", line)
        if m:
            return m.group(1).strip()
        return None

    for line in lines:
        heading = is_section_header(line)
        if heading:
            _flush_block(doc, current_heading, block)
            block = []
            current_heading = heading
            doc.add_heading(heading, level=1)
        else:
            block.append(line)

    _flush_block(doc, current_heading, block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="path to .md report file")
    ap.add_argument("--out", help="output .docx path (default: same dir, .docx extension)")
    args = ap.parse_args()

    with open(args.input, encoding="utf-8") as f:
        md_text = f.read()

    # try to extract name/symbol/isin from first heading line
    m = re.search(r"#\s+Deep Dive.*?([A-Z][^(]+)\s+\((\w+)\s*/\s*(INE\w+)\)", md_text)
    name = m.group(1).strip() if m else os.path.splitext(os.path.basename(args.input))[0]
    symbol = m.group(2) if m else ""
    isin   = m.group(3) if m else ""

    docx_bytes = md_to_docx(md_text, name, symbol, isin)
    out_path = args.out or os.path.splitext(args.input)[0] + ".docx"
    with open(out_path, "wb") as f:
        f.write(docx_bytes)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
